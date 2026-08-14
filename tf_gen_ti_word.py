from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ========== Style Setup ==========
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ========== Title ==========
title = doc.add_heading('天鯨威胁情报服务（ti_server）与流影对接说明', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run('文档版本：V1.0    更新日期：2026-08-14')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph('')

# ========== 1. Overview ==========
doc.add_heading('一、概述', level=1)
doc.add_paragraph(
    '天鯨威胁情报服务（ti_server）是一款独立的威胁情报查询与管理系统，提供 JWT 签发、'
    '情报查询、全量数据导出、客户端管理等功能。流影（Flow Shadow）平台作为其客户端，'
    '通过威胁情报配置页面向 ti_server 发起查询。'
)
doc.add_paragraph(
    '本文档详细说明 ti_server 与流影之间的认证逻辑、token 与 key 的关系，'
    '以及在 ti_server 上重置各类密钥后对流影的影响和相应操作。'
)

# ========== 2. 架构总览 ==========
doc.add_heading('二、架构总览', level=1)

doc.add_heading('2.1 部署架构', level=2)
doc.add_paragraph(
    'ti_server 采用双端口架构：'
)

items = [
    ('管理端口（默认 8090）', '提供 Web 管理界面和 API，包括客户端管理、系统配置、'
     'JWT 密钥管理、HTTPS 证书管理等'),
    ('查询端口（默认 8091）', '对外提供威胁情报查询协议，包括 JWT 签发和情报查询，'
     '仅暴露查询能力，不暴露管理功能'),
]
for label, desc in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.2 流影对接的两种服务', level=2)

# Table: 两种服务对比
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['项目', '威胁情报查询（threatinfo）', '高级情报服务（threatinfopro）']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

data = [
    ['配置文件', '/Server/etc/tisrs.conf', '/Server/etc/tic.conf'],
    ['配置字段', 'KEY / HOST / PORT', 'API_KEY / HOST / PORT'],
    ['认证方式', 'KEY 换取 JWT（向 ti_server 请求）', 'API_KEY 本地计算 MD5 token'],
    ['传输协议', 'HTTP', 'HTTPS'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph('')

# ========== 3. 认证逻辑详解 ==========
doc.add_heading('三、认证逻辑详解', level=1)

doc.add_heading('3.1 威胁情报查询（threatinfo）— JWT 代理认证', level=2)
doc.add_paragraph(
    'threatinfo 采用 JWT（JSON Web Token）代理认证模式。流影作为代理，'
    '用 KEY 向 ti_server 换取 JWT，再携带 JWT 发起情报查询。'
)

doc.add_paragraph('完整请求流程：', style='List Bullet')

# Flow diagram
steps = [
    ('Step 1', '前端用户输入 IP/域名，调用 /d/threatinfo?key=8.8.8.8'),
    ('Step 2', 'threatinfo.cpp 读取 tisrs.conf，获取 KEY/HOST/PORT'),
    ('Step 3', f'向 ti_server 查询端口发起请求：\n'
               f'POST http://{{HOST}}:{{PORT}}/apisix/plugin/jwt/sign?key={{KEY}}'),
    ('Step 4', 'ti_server 的 api_sign() 验证 KEY：\n'
               '  • 优先匹配 t_client.cli_key（校验启用状态 + 来源 IP 白名单）\n'
               '  • 兼容旧版 service_key 配置'),
    ('Step 5', '验证通过后，用 jwt_secret 签发 JWT：\n'
               f'  make_jwt({{"sub":"ti_query", "cid":客户端ID, "exp":当前时间+过期时间}}, jwt_secret)\n'
               '  响应体即 JWT 文本（纯文本格式兼容流影）'),
    ('Step 6', '流影收到 JWT，将其附加到原始查询参数：&jwt={JWT}\n'
               '  拼接 URL 后转发到 ti_server 的 /query 接口'),
    ('Step 7', 'ti_server 的 api_query() 验证 JWT：\n'
               '  • 用 jwt_secret 验证签名和过期时间\n'
               '  • 如果 JWT 中的 cid>0，额外校验客户端的启用状态和 IP 白名单'),
    ('Step 8', '验证通过后执行情报查询，返回结果到流影前端'),
]

for label, desc in steps:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('3.2 高级情报服务（threatinfopro）— HMAC 本地签名', level=2)
doc.add_paragraph(
    'threatinfopro 采用基于时间窗口的 HMAC 式本地签名模式。'
    'token 由流影根据 API_KEY 和当前时间自行计算，无需向 ti_server 请求。'
)

doc.add_paragraph('完整请求流程：', style='List Bullet')

steps2 = [
    ('Step 1', '前端调用 /d/threatinfopro'),
    ('Step 2', 'threatinfopro.cpp 读取 tic.conf，获取 API_KEY/HOST/PORT'),
    ('Step 3', f'本地计算 token（不请求外部服务）：\n'
               f'  interval_start = timestamp - (timestamp - AUTH_START_TIME) % AUTH_VALID_TIME\n'
               f'  token = MD5(API_KEY + interval_start + valid_time)\n'
               f'  其中 valid_time = AUTH_VALID_TIME（固定时长的时间窗口）'),
    ('Step 4', '将 token 附加到原始查询参数：&token={token}'),
    ('Step 5', f'转发到 https://{{HOST}}:{{PORT}}/processor/threatinfopro'),
    ('Step 6', 'ti_server 收到请求后用同样的算法验证 token（时间窗口容忍）'),
]

for label, desc in steps2:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('3.3 两种认证方式对比', level=2)

table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(['对比项', 'threatinfo（JWT）', 'threatinfopro（MD5）']):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

data2 = [
    ['凭证来源', 'KET → 请求 ti_server 获取 JWT', 'API_KEY → 本地计算 MD5'],
    ['是否网络请求', '是（每次查询请求 JWT 一次）', '否（纯本地计算）'],
    ['传输协议', 'HTTP', 'HTTPS'],
    ['时间敏感性', 'JWT 有过期时间（默认 3600s）', '时间窗口校验（可容忍偏差）'],
    ['适用场景', '标准威胁情报查询', '高级批量查询与处理'],
]
for row_idx, row_data in enumerate(data2):
    for col_idx, val in enumerate(row_data):
        table2.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph('')

# ========== 4. Token 管理机制 ==========
doc.add_heading('四、Token 管理机制', level=1)

doc.add_paragraph(
    'ti_server 管理着两个不同层级的 token，分别服务于不同场景：'
)

doc.add_heading('4.1 JWT 签名密钥（jwt_secret）', level=2)
doc.add_paragraph(
    'jwt_secret 是 ti_server 系统级别的配置项，存储在 t_config 表中（k="jwt_secret"）。'
    '它是签发和验证所有 JWT 的对称密钥。ti_server 初始化时自动随机生成 64 位十六进制字符串。'
)
doc.add_paragraph('配置位置：ti_server 管理界面 → 系统配置 → JWT 签名密钥')

doc.add_heading('4.2 客户端 Key（t_client.cli_key）', level=2)
doc.add_paragraph(
    '每个客户端在 t_client 表中有一条记录，包含 cli_key 字段。'
    '该 Key 用于调用 POST /apisix/plugin/jwt/sign 换取 JWT。'
    '对应流影的 tisrs.conf 中的 KEY 字段。'
)
doc.add_paragraph('管理位置：ti_server 管理界面 → 客户端管理 → 各客户端的 Key 列')

doc.add_heading('4.3 客户端 Token（t_client.cli_token）', level=2)
doc.add_paragraph(
    '每个客户端独立的长效 Token，用于免 JWT 流程直接查询（/query?token=xxx）'
    '或全量数据导出（/export?token=xxx）。'
    '流影默认走 JWT 路径，不直接使用该 Token。'
)
doc.add_paragraph('管理位置：ti_server 管理界面 → 客户端管理 → 各客户端的 Token 列')

doc.add_heading('4.4 JWT 过期时间（jwt_expire）', level=2)
doc.add_paragraph(
    'JWT 的过期时间（秒），存储在 t_config 表（k="jwt_expire"）。'
    '默认值：3600 秒（1 小时）。'
    '流影每次查询前重新获取 JWT，过期时间不影响流影正常工作。'
)

# ========== 5. 配置字段说明 ==========
doc.add_heading('五、配置字段说明', level=1)

doc.add_paragraph('流影威胁情报配置弹窗中各字段与 ti_server 的对应关系：')

table3 = doc.add_table(rows=7, cols=3)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(['UI 字段', '配置文件', '对应 ti_server 字段']):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

data3 = [
    ['威胁情报查询 → API Key', 'tisrs.conf → KEY', 't_client.cli_key'],
    ['威胁情报查询 → 服务地址', 'tisrs.conf → HOST', 'ti_server 查询端口 IP'],
    ['威胁情报查询 → 服务端口', 'tisrs.conf → PORT', 'ti_server 查询端口 PORT（8091）'],
    ['高级情报服务 → API Key', 'tic.conf → API_KEY', 't_client.cli_key（同）'],
    ['高级情报服务 → 服务地址', 'tic.conf → HOST', 'ti_server 查询端口 IP'],
    ['高级情报服务 → 服务端口', 'tic.conf → PORT', 'ti_server 查询端口 PORT（8091）'],
]
for row_idx, row_data in enumerate(data3):
    for col_idx, val in enumerate(row_data):
        table3.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph('')

# ========== 6. 重置场景及影响 ==========
doc.add_heading('六、重置场景及对流影的影响', level=1)

doc.add_paragraph(
    'ti_server 上可重置三类凭证，不同重置操作对流影的影响程度不同。'
)

doc.add_heading('6.1 客户端 Key 重置', level=2)
p = doc.add_paragraph()
run = p.add_run('触发位置：')
run.bold = True
p.add_run('ti_server 客户端管理 → 点击客户端的"Key"列的重新生成按钮')

p = doc.add_paragraph()
run = p.add_run('操作效果：')
run.bold = True
p.add_run(f'生成新的 cli_key（secrets.token_hex(16)，32 位十六进制），更新 t_client 表')

p = doc.add_paragraph()
run = p.add_run('影响范围：')
run.bold = True
p.add_run('流影 threatinfo 查询全部失效')

p = doc.add_paragraph()
run = p.add_run('原因分析：')
run.bold = True
p = doc.add_paragraph()
p.add_run(
    '流影每次调用 POST /apisix/plugin/jwt/sign?key=旧KEY 时，'
    'ti_server 在 t_client 表中找不到匹配的 cli_key，返回 "key invalid"。'
    '流影无法获取 JWT，因而所有威胁情报查询失败。'
)

p = doc.add_paragraph()
run = p.add_run('流影需要做的操作：')
run.bold = True
p.add_run('更新 tisrs.conf 和 tic.conf 中的 KEY/API_KEY 为新值。')

# 操作步骤
steps_key = [
    '打开流影 Web 界面 → 情报按钮 → 威胁情报服务配置',
    '将"威胁情报查询"的 API Key 字段填入新 Key',
    '将"高级情报服务"的 API Key 字段填入同一个新 Key',
    '点击"保存"按钮',
]
for i, s in enumerate(steps_key):
    doc.add_paragraph(f'{i+1}. {s}', style='List Number')

doc.add_heading('6.2 客户端 Token 重置', level=2)
p = doc.add_paragraph()
run = p.add_run('触发位置：')
run.bold = True
p.add_run('ti_server 客户端管理 → 点击客户端的"Token"列的重新生成按钮')

p = doc.add_paragraph()
run = p.add_run('操作效果：')
run.bold = True
p.add_run(f'生成新的 cli_token（secrets.token_hex(24)，48 位十六进制），更新 t_client 表')

p = doc.add_paragraph()
run = p.add_run('影响范围：')
run.bold = True
p.add_run('对流影无影响')

p = doc.add_paragraph()
run = p.add_run('原因分析：')
run.bold = True
p = doc.add_paragraph()
p.add_run(
    '流影 threatinfo 走的是 KEY → JWT → 查询路径，'
    '不直接使用 t_client.token。Token 仅用于直查场景'
    '（curl /query?token=xxx 或全量数据导出 /export?token=xxx）。'
    '流影不存储也不使用该 Token。'
)

p = doc.add_paragraph()
run = p.add_run('流影需要做的操作：')
run.bold = True
p.add_run('无')

doc.add_heading('6.3 JWT 签名密钥重置', level=2)
p = doc.add_paragraph()
run = p.add_run('触发位置：')
run.bold = True
p.add_run('ti_server 管理界面 → 系统配置 → 修改 JWT 签名密钥')

p = doc.add_paragraph()
run = p.add_run('操作效果：')
run.bold = True
p.add_run('更新 t_config 中 jwt_secret 的值，所有已签发的旧 JWT 立即失效')

p = doc.add_paragraph()
run = p.add_run('影响范围：')
run.bold = True
p.add_run('对流影无影响')

p = doc.add_paragraph()
run = p.add_run('原因分析：')
run.bold = True
p = doc.add_paragraph()
p.add_run(
    '流影每次查询前都重新调用 POST /apisix/plugin/jwt/sign 获取 JWT，'
    '不缓存旧 JWT。只要 KEY 未变，流影就能拿到新密钥签发的 JWT。'
    '切换过程无缝，用户无感知。'
)

p = doc.add_paragraph()
run = p.add_run('流影需要做的操作：')
run.bold = True
p.add_run('无')

# ========== 7. 总结对比表 ==========
doc.add_heading('七、重置影响速查表', level=1)

table4 = doc.add_table(rows=4, cols=4)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(['重置项', '对应配置', '对流影影响', '需要修改流影']):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True

data4 = [
    ['客户端 Key', 'tisrs.conf / tic.conf 的 KEY', '查询全部失效', '是 — 更新 KEY 字段'],
    ['客户端 Token', '无（流影不使用）', '无', '否'],
    ['JWT 签名密钥', '无（流影自动获取）', '无', '否'],
]
for row_idx, row_data in enumerate(data4):
    for col_idx, val in enumerate(row_data):
        table4.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph('')

# ========== 8. 常见问题 ==========
doc.add_heading('八、常见问题', level=1)

faqs = [
    ('Q1: 为什么流影 UI 没有 token 输入框？',
     'token 是流影服务端自动获取或计算的中间产物，不是用户需要管理的凭证。'
     'threatinfo 的 JWT 由流影每次查询前自动向 ti_server 换取；'
     'threatinfopro 的 MD5 token 由流影根据 API_KEY + 时间窗口本地计算。'
     '用户只需提供原始的 KEY/API_KEY，token 对用户透明。'),
    ('Q2: 同一台 ti_server 可以为多个流影服务吗？',
     '可以。ti_server 的 t_client 表支持多个客户端，每个客户端有独立的 Key、Token、'
     '启用状态、来源 IP 白名单和允许更新截止日期。每个流影实例对应一个客户端。'),
    ('Q3: 如何验证威胁情报配置是否正常？',
     '在流影威胁情报配置弹窗中点击"测试"按钮。该测试会读取当前配置的 KEY/HOST/PORT，'
     '向 ti_server 发起 JWT 签发请求，如果返回 "连通正常，JWT 获取成功" 则配置正确。'),
    ('Q4: ti_server 支持 HTTPS 吗？',
     'threatinfo 服务采用 HTTP（已在 threatinfo.cpp 中硬编码），threatinfopro 采用 HTTPS（已在 '
     'threatinfopro.cpp 中硬编码 "https://"）。ti_server 管理界面支持上传 PFX 证书启用 HTTPS。'),
    ('Q5: JWT 过期了需要手动刷新吗？',
     '不需要。流影每次查询前都重新获取 JWT，不会缓存过期 JWT。'
     '即使 JWT 过期，下一次查询时流影会重新调用 /apisix/plugin/jwt/sign 获取新的 JWT。'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    doc.add_paragraph(a)
    doc.add_paragraph('')

# ========== Save ==========
output_path = 'd:\\QorderProject\\SOC\\天鯨威胁情报服务与流影对接说明.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')